from docx import Document
from datetime import datetime
import argparse

parser = argparse.ArgumentParser()
parser.add_argument("--template")
parser.add_argument("--output")
parser.add_argument("--empleado")
parser.add_argument("--puesto")
parser.add_argument("--fecha")
args = parser.parse_args()

doc = Document(args.template)

# 📅 FECHAS
fecha_dt = datetime.strptime(args.fecha, "%d/%m/%Y")

datos = {
    "{{EMPLEADO}}": args.empleado,
    "{{PUESTO}}": args.puesto,
    "{{GERENTE}}": "Reyna Guadalupe Deantes Delgado",
    "{{FECHA}}": datetime.now().strftime("%d/%m/%Y"),
    "{{DIA}}": fecha_dt.strftime("%d"),
    "{{MES}}": fecha_dt.strftime("%B"),
    "{{ANIO}}": fecha_dt.strftime("%Y"),
}

def reemplazar_texto(doc, datos):
    for p in doc.paragraphs:
        texto_original = p.text

        for key, value in datos.items():
            if key in texto_original:
                texto_original = texto_original.replace(key, value)

        # 🔥 reescribir TODO el párrafo
        if p.text != texto_original:
            p.clear()
            p.add_run(texto_original)

    # 🔥 también en tablas
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                reemplazar_texto_en_parrafos(cell.paragraphs, datos)


def reemplazar_texto_en_parrafos(parrafos, datos):
    for p in parrafos:
        texto_original = p.text

        for key, value in datos.items():
            if key in texto_original:
                texto_original = texto_original.replace(key, value)

        if p.text != texto_original:
            p.clear()
            p.add_run(texto_original)

doc.save(args.output)
