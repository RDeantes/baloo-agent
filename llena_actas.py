from docx import Document
from datetime import datetime
import argparse


parser = argparse.ArgumentParser()
parser.add_argument("--template")
parser.add_argument("--output")
parser.add_argument("--empleado")
parser.add_argument("--puesto")
parser.add_argument("--fecha")
args = parser.parse_args()

doc = Document(args.template)

fecha_dt = datetime.strptime(args.fecha, "%d/%m/%Y")

datos = {
    "{{EMPLEADO}}": args.empleado,
    "{{PUESTO}}": args.puesto,
    "{{FECHA}}": datetime.now().strftime("%d/%m/%Y"),
    "{{DIA}}": fecha_dt.strftime("%d"),
    "{{MES}}": fecha_dt.strftime("%B"),
    "{{ANIO}}": fecha_dt.strftime("%Y"),
    "{{GERENTE}}": "Reyna Guadalupe Deantes Delgado"
}

for p in doc.paragraphs:
    for key, value in datos.items():
        if key in p.text:
            p.text = p.text.replace(key, value)

for table in doc.tables:
    for row in table.rows:
        for cell in row.cells:
            for p in cell.paragraphs:
                for key, value in datos.items():
                    if key in p.text:
                        p.text = p.text.replace(key, value)

doc.save(args.output)
